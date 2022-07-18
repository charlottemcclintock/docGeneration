
# import modules
import docx
from docx.shared import Pt
from docx.shared import Inches
import os
import glob
import pyodbc
import pandas as pd
from class_structure import Viz

class BuildDocument(): 
    '''
    Code to build Word document from visualizations and text in mission folders
    using python-docx.

    Params: 
        place (str): place entity
    
    Returns: 
        demo-{place}.docx - compiled Word document, saved in place folder.
    '''

    def __init__(self, place):

        # add mission as attribute
        self.place = place

        # run method below 
        self.getPlaceInfo()
        self.importText()
        self.buildDocument()


    def getPlaceInfo(self):
        '''
        Get basic place information and save as class attribute.
        '''
        # define query to get USS compliance by service type
        query = '''OMITTED'''

        # format sql statement and query data 
        place_query = query.format("'"+ self.place + "'")
        overview = pd.read_sql(place_query, Viz().cnxn)
        
        # save post info into class attribute
        self.overview = overview

    def importText(self):
        '''
        Import all text from standard_text and mission specific folders, 
        save to a central dictionary for easy use in compiling document.

        Returns: 
            text_sections (dict): keys are filenames of .txt files, values are 
            the text they contain. 
        '''

        # get list of paths to both specific and standard text using glob
        specific_text = glob.glob(os.path.join(os.getcwd(), "..", "test_examples", self.mission, "*.txt"))
        standard_text = glob.glob(os.path.join(os.getcwd(),  "standard_text", "*.txt"))
        all_text = specific_text + standard_text

        # create dict with filenames and text sections
        text_sections = dict()
        for file_path in all_text:
            # clean section name from file name
            section = os.path.basename(file_path)
            section = section.replace('.txt', '')
            # add each file to dictionary section
            with open(file_path) as f_input:
                text_sections.update({section: f_input.read()})

        self.text_sections = text_sections

    def buildDocument(self):

        '''
        Compile document sections and write out document. 

        Returns: 
            demo-{place}.docx - compiled Word document, saved in place folder.
        '''

        # instantiate document class
        doc = docx.Document()

        '''Define some helper functions to simplify later code '''

        # add bold line
        def add_bold_line(text):
            bold_text = doc.add_paragraph()
            bold_text.add_run(text).bold=True

        # add bold, uppercase header
        def add_header(text): 
            header = doc.add_paragraph()
            header.add_run(text.upper()).bold=True

        # add italic, uppercase subheader
        def add_subheader(text): 
            header = doc.add_paragraph()
            header.add_run(text.upper()).italic=True

        # add figure from mission folder, pass in name and width
        def add_figure(name, width):
            doc.add_picture(f'../test_examples/{self.place}/{name}.png', width=Inches(width))
            # note: if don't use if inserting image from other location

        '''General document styling'''

        # set overall fonts
        font = doc.styles['Normal'].font
        font.name = 'Times New Roman'

        '''Begin constructing main section of document.'''

        # add introduction header
        add_header('Introduction')
        # add standard introduction text, format with region/environment score
        doc.add_paragraph(self.text_sections['intro'].format(self.overview['Region'][0], self.overview['Score'][0]))

        # overview section
        add_header('Overview')
        #doc.add_paragraph(self.text_sections['overview'])
        #add_figure('overview', 5.5)

        doc.add_paragraph(self.text_sections['timeseries'])
        add_figure('timeseries', 5.5)

        # top 5/bottom 5 - icass
        add_header('Top & Bottom Services')
        doc.add_paragraph(self.text_sections['topbottom'])
        add_figure('top5bottom5', 6)

        # write out document
        doc.save(f'../test_examples/{self.place}/demo-{self.place}.docx')
        print('Wrote out document.')

# %%
if __name__ == "__main__":
    instance = BuildDocument(place='Country')

# %%
